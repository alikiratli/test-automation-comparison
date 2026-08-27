"""Karsilastirma raporunu (.docx) uretir.

Kullanim:
    pip install python-docx
    python docs/generate_report.py

Cikti: Otomasyon_Karsilastirma_Raporu.docx (proje kokunde)

NOT: Rapordaki tum sayisal degerler bu makinede GERCEKTEN kosturulan
testlerden alinmistir; tahmin degildir.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Otomasyon_Karsilastirma_Raporu.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x59, 0x59, 0x59)
GOOD = RGBColor(0x1E, 0x7A, 0x3C)
BAD = RGBColor(0xB3, 0x1B, 0x1B)


# --------------------------------------------------------------------------- #
# Bicimlendirme yardimcilari
# --------------------------------------------------------------------------- #
def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15


def h1(doc: Document, text: str) -> None:
    doc.add_page_break()
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = ACCENT


def h2(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = ACCENT


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def para(doc: Document, text: str, bold: bool = False, italic: bool = False,
         color: RGBColor | None = None, size: float | None = None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)


def bullets(doc: Document, items: list[str], style: str = "List Bullet") -> None:
    for item in items:
        doc.add_paragraph(item, style=style)


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def table(doc: Document, headers: list[str], rows: list[list[str]],
          widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = t.rows[0].cells
    for index, title in enumerate(headers):
        header_cells[index].text = ""
        run = header_cells[index].paragraphs[0].add_run(title)
        run.bold = True
        run.font.size = Pt(9.5)

    for row in rows:
        cells = t.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = ""
            run = cells[index].paragraphs[0].add_run(str(value))
            run.font.size = Pt(9.5)

    if widths:
        for row in t.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)

    doc.add_paragraph()


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(12)


# --------------------------------------------------------------------------- #
# Rapor icerigi
# --------------------------------------------------------------------------- #
def build_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TEST OTOMASYON ARAÇLARI\nKARŞILAŞTIRMA RAPORU")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Selenium WebDriver  ·  Robot Framework  ·  Playwright")
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Aynı uygulama, aynı test senaryoları, üç farklı teknoloji\n"
        "Uygulanan ve gerçekten çalıştırılan üç otomasyon projesinin analizi"
    )
    run.italic = True
    run.font.size = Pt(11)

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Test Uygulaması: SauceDemo (https://www.saucedemo.com)")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Toplam 183 test  ·  7.002 satır kod  ·  Windows 10 / Python 3.13 / Chrome 151")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED


def build_toc(doc: Document) -> None:
    doc.add_page_break()
    h2(doc, "İçindekiler")
    entries = [
        "1. Yönetici Özeti",
        "2. Projenin Kurgusu: Neden Aynı Uygulama?",
        "3. Test Edilen Uygulama ve Senaryolar",
        "4. Proje 1 — Selenium WebDriver + pytest",
        "5. Proje 2 — Robot Framework + SeleniumLibrary",
        "6. Proje 3 — Playwright + pytest",
        "7. Ölçülen Sonuçlar",
        "8. Ayrıntılı Karşılaştırma",
        "9. Uygulama Sırasında Karşılaşılan Gerçek Sorunlar",
        "10. Hangi Durumda Hangisi? — Karar Rehberi",
        "11. Sonuç ve Öneri",
        "Ek A — Çalıştırma Komutları",
        "Ek B — Dosya Yapısı",
    ]
    for entry in entries:
        p = doc.add_paragraph(entry)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)


def section_1(doc: Document) -> None:
    h1(doc, "1. Yönetici Özeti")

    para(doc,
         "Bu rapor, aynı web uygulamasını (SauceDemo) aynı test senaryolarıyla test eden "
         "üç ayrı otomasyon projesinin uygulanmasına ve gerçekten çalıştırılmasına dayanır. "
         "Amaç, üç teknolojiyi teorik olarak değil, aynı işi yaptırarak karşılaştırmaktır.")

    h2(doc, "En kısa cevap")
    table(doc,
          ["Teknoloji", "Bir cümlede", "Kime uygun"],
          [
              ["Selenium + pytest",
               "En yaygın, en esnek, en çok altyapı yazdıran seçenek.",
               "Kurumsal ortamlar, geniş tarayıcı/cihaz matrisi, mevcut Selenium birikimi"],
              ["Robot Framework",
               "En okunabilir ve en iyi raporlayan, programlama gerektirmeyen seçenek.",
               "Karma ekipler, iş analisti katkısı, kabul testleri, regülasyonlu sektörler"],
              ["Playwright + pytest",
               "En hızlı, en kararlı, en çok yeteneği hazır getiren seçenek.",
               "Yeni kurulan projeler, modern SPA'lar, hız ve kararlılık önceliği"],
          ],
          widths=[1.5, 2.6, 2.4])

    h2(doc, "Ölçülen üç temel bulgu")

    para(doc, "1) Hız farkı büyük ve gerçek.", bold=True)
    para(doc,
         "Aynı 15–16 login senaryosu Selenium ile 77 saniye, Playwright ile 24 saniye sürdü "
         "(≈3,2 kat). Tüm pakette Selenium'un 61 testi 6 dakika 7 saniye, Playwright'ın "
         "75 testi 1 dakika 37 saniye sürdü. Fark, Playwright'ın her test için yeni tarayıcı "
         "süreci yerine yeni tarayıcı bağlamı (context) açmasından ve WebDriver ara katmanının "
         "olmamasından kaynaklanır.")

    para(doc, "2) Kararlılık farkı, kod miktarı farkından daha önemli.", bold=True)
    para(doc,
         "Selenium ve Robot projelerinde toplam 8 gerçek hatanın 5'i doğrudan bekleme/tıklama "
         "kararlılığıyla ilgiliydi. Playwright projesinde bu sınıftan hiç hata çıkmadı. "
         "Selenium tarafında kararlılık için yazılan altyapı (özel bekleme koşulları, yeniden "
         "deneme mekanizmaları) 415 satır tuttu; Playwright'ta bu katmanın karşılığı yok.")

    para(doc, "3) Robot Framework'ün üstünlüğü hızda değil, okunabilirlik ve raporda.", bold=True)
    para(doc,
         "Robot'un ürettiği log.html, hiçbir kod yazılmadan her adımı, argümanını ve süresini "
         "ağaç yapısında gösterir. Test dosyaları, teknik olmayan bir ekip üyesinin "
         "okuyabileceği kadar sadedir. Buna karşılık hesaplama ve veri işleme gerektiren her "
         "iş için Python'a inmek gerekir — bu projede 200 satırlık ayrı bir kütüphane yazıldı.")

    h2(doc, "Öneri")
    para(doc,
         "Yeni bir web otomasyon projesine bugün başlanıyorsa Playwright öneriliyor. "
         "Ekipte programlama bilmeyen üyeler test yazacaksa veya raporun kurum içinde "
         "kanıt olarak kullanılması gerekiyorsa Robot Framework öneriliyor — bu durumda "
         "SeleniumLibrary yerine Robot'un Browser (Playwright tabanlı) kütüphanesi ile "
         "her iki avantaj birleştirilebilir. Selenium, mevcut yatırımın korunması, çok geniş "
         "tarayıcı desteği veya Appium ile mobil entegrasyon gerektiğinde tercih edilmelidir.",
         )


def section_2(doc: Document) -> None:
    h1(doc, "2. Projenin Kurgusu: Neden Aynı Uygulama?")

    para(doc,
         "Framework karşılaştırmalarının çoğu elma ile armudu kıyaslar: farklı sitelerde, "
         "farklı senaryolarla, farklı olgunlukta yazılmış kodlara bakılır. Böyle bir "
         "karşılaştırmadan çıkan sonuç, framework'ün değil yazarın becerisinin ölçüsüdür.")

    para(doc, "Bu projede tek değişken framework olacak şekilde kurgu yapıldı:", bold=True)
    bullets(doc, [
        "Aynı uygulama: SauceDemo (https://www.saucedemo.com)",
        "Aynı senaryolar: login, ürün listesi, sıralama, sepet, checkout, uçtan uca akış",
        "Aynı test verisi: üç proje de ../shared/users.json ve products.json dosyalarını okur",
        "Aynı dil: üçü de Python — dil farkının karşılaştırmayı kirletmemesi için",
        "Aynı mimari olgunluk: üçü de katmanlı (test / sayfa / çekirdek), sabit kodlanmış "
        "değer içermeyen, yapılandırılabilir projeler",
    ])

    h2(doc, "Neden üçü de Python?")
    para(doc,
         "Playwright'ın en yaygın kullanımı TypeScript'tir. Buna rağmen bu projede Python "
         "seçildi; çünkü TypeScript ile yazılsaydı görülen farkların bir bölümü "
         "\"TypeScript vs Python\" farkı olurdu, framework farkı değil. Bu tercihin bedeli, "
         "Playwright'ın Python sürümünde bulunmayan bazı özelliklerin "
         "(expect.soft, toHaveScreenshot) raporda ayrıca not edilmesidir.")

    h2(doc, "Ortak test verisi")
    code_block(doc,
               "Automation/\n"
               "├── shared/\n"
               "│   ├── users.json      <- 8 login senaryosu (3 proje de okur)\n"
               "│   └── products.json   <- 6 ürün, fiyatlar, KDV oranı, müşteri bilgisi\n"
               "├── 01-selenium-pytest/\n"
               "├── 02-robot-framework/\n"
               "└── 03-playwright-pytest/")
    caption(doc,
            "Test verisinin tek kaynaktan gelmesi, üç projenin gerçekten aynı şeyi test "
            "ettiğini garanti eder. Yeni bir kullanıcı senaryosu eklendiğinde üç proje de "
            "onu otomatik olarak kapsar.")


def section_3(doc: Document) -> None:
    h1(doc, "3. Test Edilen Uygulama ve Senaryolar")

    para(doc,
         "SauceDemo, Sauce Labs tarafından test otomasyonu eğitimi için yayınlanan bir "
         "e-ticaret demo uygulamasıdır. Seçilme nedenleri: herkese açık ve kararlı olması, "
         "React ile yazılmış gerçek bir SPA olması ve kasıtlı olarak bozulmuş kullanıcı "
         "hesapları içermesi.")

    h2(doc, "Test hesapları")
    table(doc,
          ["Kullanıcı", "Davranış", "Neyi test eder"],
          [
              ["standard_user", "Normal akış", "Mutlu yol (happy path)"],
              ["locked_out_user", "Giriş engellenir", "Yetkilendirme iş kuralı"],
              ["problem_user", "Bozuk görseller, hatalı form alanları", "Bilinen kusurların izlenmesi"],
              ["performance_glitch_user", "Kasıtlı gecikme", "Bekleme stratejisinin dayanıklılığı"],
          ],
          widths=[1.8, 2.2, 2.5])

    h2(doc, "Kapsanan senaryolar")
    table(doc,
          ["#", "Senaryo", "Selenium", "Robot", "Playwright"],
          [
              ["1", "Login — geçerli/geçersiz/kilitli (veri odaklı, 8 senaryo)", "✔", "✔", "✔"],
              ["2", "Ürün listesi — sayı, isim ve fiyat doğrulaması", "✔", "✔", "✔"],
              ["3", "Sıralama — A-Z, Z-A, fiyat artan/azalan", "✔", "✔", "✔"],
              ["4", "Sepet — ekle/çıkar/rozet/kalıcılık", "✔", "✔", "✔"],
              ["5", "Checkout — form validasyonu (4 senaryo)", "✔", "✔", "✔"],
              ["6", "Checkout — KDV ve toplam matematiği", "✔", "✔", "✔"],
              ["7", "Uçtan uca satın alma akışı", "✔", "✔", "✔"],
              ["8", "Navigasyon, menü, çıkış, oturum güvenliği", "✔", "✔", "✔"],
              ["9", "Konsol hatalarının izlenmesi", "Kısmi", "Kısmi", "✔"],
              ["10", "Ağ katmanı mock'lama, offline mod", "✖", "✖", "✔"],
              ["11", "Görsel regresyon, mobil emülasyon", "✖", "✖", "✔"],
              ["12", "API testleri (aynı framework içinde)", "✖", "✖", "✔"],
              ["13", "Erişilebilirlik kontrolleri", "✖", "✖", "✔"],
          ],
          widths=[0.35, 3.2, 1.0, 0.8, 1.1])
    caption(doc,
            "1–8 arası senaryolar üç projede de birebir aynıdır ve karşılaştırmanın adil "
            "zeminini oluşturur. 9–13 arası, Playwright'ın diğer ikisinde karşılığı olmayan "
            "yeteneklerini göstermek için bilinçli olarak eklenmiştir.")

    h2(doc, "İş kuralı testinin önemi")
    para(doc,
         "Üç projede de en değerli test, \"butona tıklandı mı\" değil \"hesap doğru mu\" "
         "sorusunu soran testtir. Checkout özet sayfasında şu doğrulama yapılır:")
    code_block(doc,
               "KDV        = Ara toplam × 0,08\n"
               "Genel toplam = Ara toplam + KDV\n"
               "Ara toplam  = Sepetteki ürün fiyatlarının toplamı (referans veriyle karşılaştırılır)")
    para(doc,
         "Bu doğrulama üç projede de farklı yerlerde yaşar: Selenium'da bir dataclass "
         "metodunda (CheckoutTotals.validate), Robot'ta bir Python keyword'ünde "
         "(Verify Checkout Totals), Playwright'ta yine bir dataclass'ta. Bu, "
         "\"iş mantığı UI katmanından ayrılmalı\" ilkesinin üç framework'te de geçerli "
         "olduğunu gösterir.")


def section_4(doc: Document) -> None:
    h1(doc, "4. Proje 1 — Selenium WebDriver + pytest")

    para(doc, "Klasör: 01-selenium-pytest/  ·  27 dosya  ·  2.748 satır  ·  61 test",
         bold=True, color=MUTED)

    h2(doc, "Ne yaptı?")
    para(doc,
         "Page Object Model (POM) deseniyle katmanlı bir otomasyon çerçevesi kuruldu. "
         "Selenium sadece tarayıcıyı sürer; test koşucusu, raporlama, paralellik, "
         "yapılandırma, loglama — hepsi ayrı ayrı seçilip birbirine bağlandı.")

    h2(doc, "Mimari")
    code_block(doc,
               "tests/       →  pages/       →  core/        →  selenium\n"
               "(doğrulama)     (davranış)      (altyapı)       (protokol)\n"
               "\n"
               "Kural 1: Test dosyaları hiçbir zaman CSS selector veya WebDriverWait görmez.\n"
               "Kural 2: Page Object'ler hiçbir zaman assert yazmaz — durum döner.\n"
               "Kural 3: core/ katmanı uygulamayı bilmez, genel amaçlıdır.")

    table(doc,
          ["Katman", "Dosyalar", "Sorumluluk"],
          [
              ["core/", "driver_factory, base_page, waits, soft_assert, exceptions, logger",
               "Tarayıcı yaşam döngüsü, bekleme, yeniden deneme, loglama"],
              ["pages/", "login, inventory, cart, checkout (3 sınıf), product_detail, "
                         "components/header",
               "Uygulama davranışını metotlara çevirir"],
              ["config/", "config.yaml, settings.py", "Ortam, tarayıcı, timeout ayarları"],
              ["utils/", "parsers, data_loader", "Metin→veri dönüşümü, ortak veri okuma"],
              ["tests/", "6 dosya, 61 test", "Yalnızca doğrulama"],
          ],
          widths=[0.9, 2.6, 2.9])

    h2(doc, "Öne çıkan teknik çözümler")

    h3(doc, "a) Implicit wait kapatıldı, sadece explicit wait kullanıldı")
    para(doc,
         "config.yaml içinde implicit_wait değeri 0'dır. Sebep: WebDriver "
         "spesifikasyonunda implicit ve explicit wait'in birlikte kullanımı tanımsız "
         "davranıştır; süreler toplanarak testleri hem yavaşlatır hem öngörülemez kılar.")

    h3(doc, "b) Özel bekleme koşulları yazıldı (core/waits.py — 127 satır)")
    para(doc,
         "Selenium'un hazır Expected Condition'ları yetmediği için 7 adet özel koşul "
         "yazıldı. En kritik olanı, elemanın konumunun iki ardışık ölçümde aynı kalmasını "
         "bekleyen element_to_be_clickable_and_stable koşuludur — animasyonlu menüde "
         "tıklamanın boşa gitmesini bu koşul engeller.")
    code_block(doc,
               "def element_to_be_clickable_and_stable(locator):\n"
               "    previous_position = {}\n"
               "    def _predicate(driver):\n"
               "        element = EC.element_to_be_clickable(locator)(driver)\n"
               "        location = (element.location['x'], element.location['y'])\n"
               "        if previous_position.get('last') == location:\n"
               "            return element          # konum sabitlendi → animasyon bitti\n"
               "        previous_position['last'] = location\n"
               "        return False\n"
               "    return _predicate")

    h3(doc, "c) StaleElementReferenceException'a karşı otomatik yeniden deneme")
    para(doc,
         "React sayfayı yeniden çizdiğinde elde tutulan eleman referansları geçersizleşir. "
         "BasePage._with_stale_retry, her etkileşimi 3 kez artan bekleme ile tekrar dener. "
         "Bu mekanizma Playwright'ta gereksizdir; locator'lar tembel (lazy) olduğu için "
         "bu istisna sınıfı orada hiç oluşmaz.")

    h3(doc, "d) Hata anında teşhis toplama")
    para(doc,
         "conftest.py içindeki pytest_runtest_makereport hook'u, bir test kaldığında "
         "ekran görüntüsü, sayfa HTML'i ve tarayıcı konsol loglarını kaydeder ve "
         "HTML rapora gömer. Bu 60 satırlık kod, Robot Framework'te tek bir ayardır "
         "(run_on_failure=Capture Page Screenshot).")

    h2(doc, "Güçlü yanları"); bullets(doc, [
        "W3C standardı — her tarayıcı, her programlama dili, her grid altyapısı destekler.",
        "Ekosistem en geniş olan: Selenium Grid, BrowserStack, Sauce Labs, Appium (mobil).",
        "Kurumsal ortamlarda kabul görmüş teknoloji; iş gücü bulmak en kolay olan.",
        "Tam kontrol: her katman istendiği gibi tasarlanır, hiçbir şey dayatılmaz.",
    ])

    h2(doc, "Zayıf yanları"); bullets(doc, [
        "core/ klasöründeki 415 satırlık altyapı pratikte her projede yeniden yazılır.",
        "Otomatik bekleme yoktur; kararlılık tamamen geliştiricinin disiplinine bağlıdır.",
        "Ağ katmanına erişimi yoktur — istek engelleme/taklit etme mümkün değildir.",
        "Test başına yeni tarayıcı süreci gerekir; izolasyon pahalıdır (~1–2 sn/test).",
        "Konsol logları yalnızca Chromium tabanlı tarayıcılarda ve sonradan çekilerek okunur.",
        "Tarayıcı sürümüne bağımlıdır: bu projede Chrome 151'in bir davranış değişikliği "
        "tüm paketi kırdı (bkz. Bölüm 9).",
    ])


def section_5(doc: Document) -> None:
    h1(doc, "5. Proje 2 — Robot Framework + SeleniumLibrary")

    para(doc, "Klasör: 02-robot-framework/  ·  14 dosya  ·  2.068 satır  ·  51 test",
         bold=True, color=MUTED)

    para(doc,
         "Önemli not: Robot Framework, Selenium'un alternatifi değildir. SeleniumLibrary'nin "
         "altında yine Selenium WebDriver çalışır. Robot Framework, Selenium'un üzerine "
         "kurulan bir soyutlama ve orkestrasyon katmanıdır. Bu yüzden Selenium'un tüm "
         "sınırları (auto-wait yokluğu, ağ erişimi olmaması) burada da geçerlidir.",
         italic=True)

    h2(doc, "Ne yaptı?")
    para(doc,
         "Aynı senaryolar keyword-driven yaklaşımla yazıldı. Robot'ta sınıf yoktur; "
         "kapsülleme dosya düzeyindedir. Her sayfa için bir .resource dosyası açılır: "
         "locator'lar *** Variables *** bölümünde, davranışlar *** Keywords *** bölümünde "
         "durur.")

    h2(doc, "Selenium/pytest karşılıkları")
    table(doc,
          ["Selenium + pytest", "Robot Framework"],
          [
              ["class LoginPage:", "resources/pages/login_page.resource"],
              ["USERNAME_INPUT = (By.CSS, ...)", "${LOGIN_USERNAME_INPUT}  css:[data-test='username']"],
              ["def login(self, ...):", "Login As  (keyword)"],
              ["conftest.py fixture", "Test Setup / Suite Setup"],
              ["@pytest.mark.parametrize", "[Template] veya DataDriver + CSV"],
              ["pytest -m smoke", "robot --include smoke"],
              ["SoftAssert sınıfı (60 satır)", "Run Keyword And Continue On Failure (hazır)"],
              ["pytest_runtest_makereport hook (60 satır)", "run_on_failure=Capture Page Screenshot"],
          ],
          widths=[3.0, 3.4])

    h2(doc, "Öne çıkan teknik çözümler")

    h3(doc, "a) [Template] ile veri odaklı test")
    para(doc,
         "Aşağıdaki tablo, raporda 5 ayrı test satırı olarak görünür. Bu tabloyu okumak "
         "veya yeni satır eklemek için Python bilmek gerekmez:")
    code_block(doc,
               "Gecersiz giris senaryolari\n"
               "    [Template]    Login Should Fail With Error\n"
               "    locked_out_user  secret_sauce    Epic sadface: Sorry, this user has been ...\n"
               "    standard_user    wrong_password  Epic sadface: Username and password do ...\n"
               "    ${EMPTY}         secret_sauce    Epic sadface: Username is required\n"
               "    standard_user    ${EMPTY}        Epic sadface: Password is required")

    h3(doc, "b) DataDriver ile CSV'den test üretimi")
    para(doc,
         "01b_login_datadriven.robot dosyasında tek bir test tanımı vardır. DataDriver "
         "kütüphanesi koşum anında data/login_scenarios.csv dosyasını okur ve her satır "
         "için ayrı bir test üretir. Bir iş analisti, tek satır kod yazmadan Excel/CSV "
         "üzerinden yeni test ekleyebilir. Bu, Robot Framework'ün en ayırt edici özelliğidir.")

    h3(doc, "c) Python kütüphanesi ile hesaplama (libraries/SauceDemoLibrary.py)")
    para(doc,
         "Robot, iş akışını tarif etmekte güçlü, hesap yapmakta zayıftır. \"$29.99 metnini "
         "sayıya çevir, KDV'yi hesapla, liste sıralı mı kontrol et\" gibi işler Robot "
         "söz diziminde okunaksızlaşır. Bu mantık 200 satırlık bir Python kütüphanesine "
         "taşındı ve Robot'a keyword olarak sunuldu:")
    code_block(doc,
               "Python                          →  Robot keyword\n"
               "parse_price()                   →  Parse Price\n"
               "verify_checkout_totals()        →  Verify Checkout Totals\n"
               "verify_list_is_sorted_ascending →  Verify List Is Sorted Ascending")

    h3(doc, "d) İş seviyesinde tek keyword")
    para(doc,
         "Uçtan uca satın alma akışının tamamı tek satıra indirildi. Teknik olmayan bir "
         "ekip üyesi bu satırı okuyup ne yapıldığını anlayabilir:")
    code_block(doc,
               "Complete Purchase For    standard_user    Sauce Labs Backpack    Sauce Labs Onesie")

    h2(doc, "Güçlü yanları"); bullets(doc, [
        "Rapor kalitesi rakipsiz: log.html her keyword'ü, argümanını ve süresini ağaç "
        "yapısında gösterir — hiçbir şey kodlamadan.",
        "Öğrenme eğrisi çok düşük; test yazmak için programlama bilmek gerekmez.",
        "Kurulum/temizlik, etiketleme, veri odaklı test dilin içinde — kütüphane değil.",
        "Aynı test dosyaları SeleniumLibrary yerine Browser (Playwright) kütüphanesine "
        "geçirilebilir: testler değil, keyword katmanı değişir.",
        "İş analisti ve manuel testçi ekiplerinin gerçekten katkı verebildiği tek yaklaşım.",
        "Denetim/regülasyon ortamlarında rapor doğrudan kanıt olarak kullanılabilir.",
    ])

    h2(doc, "Zayıf yanları"); bullets(doc, [
        "Karmaşık mantık (döngü içinde koşul, veri dönüşümü) çabuk okunaksızlaşır; "
        "Python'a inmek gerekir — libraries/ klasörünün varlık sebebi budur.",
        "Boşluk duyarlı söz dizimi (2+ boşluk ayırıcı) yeni başlayanı yorar.",
        "IDE desteği ve yeniden düzenleme (refactoring) araçları Python/TypeScript kadar "
        "olgun değildir.",
        "Sınıf ve kalıtım olmadığı için kod tekrarını engellemek daha zordur.",
        "Değişken adları büyük/küçük harf ve boşluk duyarsızdır; bu projede bir vararg "
        "global bir sözlüğü gölgeleyerek gerçek bir hataya yol açtı (bkz. Bölüm 9).",
        "SeleniumLibrary üzerinden çalıştığı için Selenium'un tüm sınırları geçerlidir.",
    ])


def section_6(doc: Document) -> None:
    h1(doc, "6. Proje 3 — Playwright + pytest")

    para(doc, "Klasör: 03-playwright-pytest/  ·  27 dosya  ·  2.361 satır  ·  75 test",
         bold=True, color=MUTED)

    h2(doc, "Ne yaptı?")
    para(doc,
         "Aynı senaryolar Playwright'ın locator tabanlı Page Object modeliyle yazıldı. "
         "Ayrıca, diğer iki framework'te karşılığı olmayan yetenekler için ayrı bir test "
         "klasörü (tests/advanced/) eklendi: ağ katmanı müdahalesi, görsel regresyon, "
         "cihaz emülasyonu, API testleri ve erişilebilirlik kontrolleri.")

    h2(doc, "Mimari farkın kanıtı: aynı işi yapan iki dosya")
    table(doc,
          ["Dosya", "Selenium", "Playwright", "Fark"],
          [
              ["core/base_page.py", "288 satır", "109 satır", "%62 daha az"],
              ["core/waits.py", "127 satır", "yok", "Dosyanın tamamı gereksiz"],
              ["Toplam altyapı", "415 satır", "109 satır", "%74 daha az"],
          ],
          widths=[1.8, 1.3, 1.3, 1.6])
    caption(doc,
            "Playwright'ın conftest.py dosyası (222 satır) Selenium'unkinden (201 satır) "
            "uzun görünür. Ancak boş satırlar, yorumlar ve açıklama blokları çıkarıldığında "
            "işlevsel kod Selenium'da 114, Playwright'ta 92 satırdır — üstelik Selenium "
            "sürümünün 60 satırı yalnızca hata anında ekran görüntüsü/HTML/konsol toplamak "
            "içindir; Playwright'ta bunun karşılığı üç komut satırı bayrağıdır.")

    h2(doc, "Farkın üç kaynağı")

    h3(doc, "a) Auto-waiting (otomatik bekleme)")
    para(doc,
         "locator.click() çağrıldığında Playwright, elemanın DOM'a eklenmesini, görünür "
         "olmasını, konumunun sabitlenmesini (animasyonun bitmesini), etkin olmasını ve "
         "üzerinde başka eleman bulunmamasını kendiliğinden bekler. Selenium'da bunların "
         "hepsi elle yazılır — core/waits.py dosyasının tamamı budur.")

    h3(doc, "b) Locator'lar tembeldir (lazy)")
    para(doc,
         "page.locator(...) DOM'da arama yapmaz; yalnızca bir tarif tutar. Arama her "
         "etkileşimde yeniden yapılır. Bu nedenle Playwright'ta "
         "StaleElementReferenceException diye bir şey yoktur — Selenium projesindeki "
         "yeniden deneme mekanizmasının tamamı burada gereksizdir.")

    h3(doc, "c) Web-first assertion'lar")
    para(doc,
         "expect(locator).to_have_text(\"X\") koşul sağlanana kadar tekrar dener. "
         "Selenium'daki assert element.text == \"X\" tek seferliktir; bu yüzden ondan önce "
         "ayrıca bir WebDriverWait yazmak gerekir — iki ayrı adım, iki ayrı hata kaynağı.")

    h2(doc, "BrowserContext: izolasyonun ucuzlaması")
    para(doc,
         "Playwright'ın en önemli mimari kararı BrowserContext'tir: izole bir tarayıcı "
         "profili (çerez, localStorage, izinler ayrı). Tarayıcı süreç olarak bir kez açılır, "
         "her test kendi context'ini alır.")
    table(doc,
          ["", "Selenium", "Playwright"],
          [
              ["İzolasyon birimi", "Tarayıcı süreci", "BrowserContext"],
              ["Test başına maliyet", "~1–2 saniye", "~20–50 milisaniye"],
              ["İzolasyon kalitesi", "Tam", "Tam"],
              ["Paralellikte bellek", "Süreç başına tam tarayıcı", "Tek tarayıcı, N context"],
          ],
          widths=[2.0, 2.2, 2.2])

    h2(doc, "Oturum yeniden kullanımı (storage_state)")
    para(doc,
         "Playwright, çerezleri ve localStorage'ı JSON olarak dışa aktarıp yeni bir "
         "context'e yükleyebilir. 50 test için 50 kez arayüzden giriş yapmak yerine bir kez "
         "giriş yapılıp aynı oturum 50 context'e enjekte edilir. Büyük paketlerde bu tek "
         "özellik toplam süreyi %30–50 kısaltabilir. Selenium'da benzeri, çerezleri elle "
         "yerleştirerek taklit edilebilir ama localStorage için JavaScript gerekir.")

    h2(doc, "Diğer ikisinde karşılığı olmayan yetenekler")
    table(doc,
          ["Yetenek", "Ne işe yarar", "Selenium/Robot'ta"],
          [
              ["page.route()", "İsteği engelle, 500 hatası döndür, sahte yanıt ver",
               "Yok — harici proxy (mitmproxy) gerekir"],
              ["set_offline()", "Çevrimdışı davranışı test et", "Yok"],
              ["Trace Viewer", "Her adımın DOM anlık görüntüsü, ağ, konsol — geri sarılabilir",
               "Yok — sadece ekran görüntüsü"],
              ["APIRequestContext", "Aynı framework içinde API testi, oturum paylaşımı",
               "Ayrı kütüphane (requests / RequestsLibrary)"],
              ["devices[...]", "~130 hazır cihaz profili (viewport + UA + dokunmatik)",
               "Sadece pencere boyutu; gerçeği için Appium"],
              ["get_by_role()", "Erişilebilirlik ağacı üzerinden seçim", "Yok"],
              ["page.on('console')", "Anlık konsol dinleme, tüm tarayıcılarda",
               "Sadece Chromium, sonradan çekilerek"],
              ["codegen / --ui", "Tıkla-kod üret; etkileşimli koşum modu", "Yok"],
          ],
          widths=[1.5, 2.6, 2.3])

    h2(doc, "Güçlü yanları"); bullets(doc, [
        "Auto-wait sayesinde kararsız (flaky) test oranı belirgin şekilde düşüktür.",
        "Trace Viewer: bir hatayı zaman içinde geri sararak incelemek — hata ayıklamada "
        "diğer ikisinin sunmadığı bir yetenek.",
        "Ağ mock'lama, cihaz emülasyonu, çoklu context, API testi — hepsi tek pakette.",
        "Context modeli sayesinde izolasyon hem tam hem ucuz; paralellik çok verimli.",
        "Tarayıcı binary'lerini kendisi indirir ve sürümü kilitler — sürüm uyumsuzluğu yaşanmaz.",
    ])

    h2(doc, "Zayıf yanları"); bullets(doc, [
        "Yalnızca Chromium/Firefox/WebKit; gerçek Chrome/Edge yerine paketlenmiş sürümler "
        "kullanılır (kanal seçilebilir, ancak IE veya macOS Safari desteği yoktur).",
        "Ekosistem Selenium kadar geniş değildir; kurumsal araç entegrasyonları daha yenidir.",
        "Python sürümü, TypeScript sürümünün gerisindedir (expect.soft, toHaveScreenshot yok).",
        "Mobil için gerçek cihaz testi yoktur — emülasyon vardır, Appium'un yerini tutmaz.",
        "Tarayıcı binary'leri yaklaşık 300 MB disk alanı ister.",
    ])


def section_7(doc: Document) -> None:
    h1(doc, "7. Ölçülen Sonuçlar")

    para(doc,
         "Aşağıdaki tüm değerler bu makinede gerçekten çalıştırılan testlerden alınmıştır. "
         "Ortam: Windows 10 Pro, Python 3.13.7, Chrome 151, headless mod, tek işlem "
         "(paralellik kapalı), aynı internet bağlantısı.", italic=True)

    para(doc,
         "Ölçüm sırasında öğrenilen önemli bir nokta: iki test paketi aynı anda "
         "çalıştırıldığında Selenium tabanlı paketlerde (Selenium ve Robot) tıklamalar "
         "toplu hâlde etkisiz kalmaya başladı ve onlarca test düştü. Aynı testler tek "
         "başına çalıştırıldığında ilk denemede geçti. Playwright paketi aynı çekişme "
         "altında etkilenmedi. Bu, kararlılık farkının yalnızca kod kalitesinden değil, "
         "kaynak baskısı altındaki davranıştan da kaynaklandığını gösterir — CI "
         "sunucularında bu koşul kuraldır, istisna değil.",
         italic=True, color=MUTED)

    h2(doc, "Koşum sonuçları")
    table(doc,
          ["Proje", "Test sayısı", "Süre", "Test başına", "Sonuç"],
          [
              ["Selenium + pytest", "61", "6 dk 07 sn (367 sn)", "6,0 sn", "61 geçti"],
              ["Robot Framework", "51", "≈9,5 dk (577 sn)", "≈11,3 sn", "50 geçti, 1 tekrar denemede"],
              ["Playwright + pytest", "75", "1 dk 37 sn (97 sn)", "1,3 sn", "74 geçti, 1 xfail"],
          ],
          widths=[1.7, 1.0, 1.7, 1.1, 1.4])
    caption(doc,
            "Not: Test sayıları birebir aynı değildir. Playwright'ta 13 test, diğer ikisinde "
            "karşılığı olmayan yeteneklere aittir. Robot'ta bir test [Template] ile birden "
            "çok satır çalıştırdığı için sayı düşük görünür. Adil karşılaştırma için aşağıdaki "
            "birebir aynı senaryo ölçümüne bakınız.")

    h2(doc, "Birebir aynı senaryo: login test dosyası")
    table(doc,
          ["Proje", "Test", "Süre", "Oran"],
          [
              ["Selenium + pytest", "16", "77 sn", "3,2×"],
              ["Robot Framework", "10", "≈115 sn", "4,8×"],
              ["Playwright + pytest", "15", "24 sn", "1,0× (referans)"],
          ],
          widths=[2.0, 0.9, 1.2, 1.6])
    caption(doc,
            "Aynı uygulama, aynı senaryolar, aynı makine. Playwright referans alınmıştır. "
            "Robot'un ek yavaşlığı, SeleniumLibrary'nin her keyword için ayrı WebDriver "
            "çağrısı yapmasından ve bu ortamda gereken etki-doğrulamalı tıklama "
            "mekanizmasından kaynaklanır.")

    h2(doc, "Kod metrikleri")
    table(doc,
          ["Ölçüt", "Selenium", "Robot", "Playwright"],
          [
              ["Toplam dosya", "27", "14", "27"],
              ["Toplam satır", "2.748", "2.068", "2.361"],
              ["Altyapı katmanı (core/)", "415 satır", "335 satır (common.resource)", "109 satır"],
              ["Bekleme kodu", "127 satır ayrı dosya", "keyword'lere dağılmış", "0 satır"],
              ["Test dosyası satırı", "≈900", "≈700", "≈1.000"],
              ["Sabit kodlanmış bekleme (sleep)", "0", "1 (örnekleme aralığı)", "0"],
          ],
          widths=[2.0, 1.4, 1.9, 1.2])

    h2(doc, "Kararlılık: ilk koşumda çıkan gerçek hatalar")
    table(doc,
          ["Proje", "Hata", "Kaç tanesi bekleme/tıklama kaynaklı"],
          [
              ["Selenium", "3", "3 (menü animasyonu, konsol gürültüsü)"],
              ["Robot Framework", "5 tur, toplam 9 test", "2 sınıf hata (tıklama, React input)"],
              ["Playwright", "1", "0 (uygulamada gerçek erişilebilirlik kusuru)"],
          ],
          widths=[1.5, 1.0, 4.0])
    caption(doc,
            "Playwright'ta çıkan tek hata framework kaynaklı değildi: erişilebilirlik "
            "testi, SauceDemo'nun envanter sayfasında hiç başlık (heading) elemanı "
            "olmadığını buldu. Bu, otomasyonun gerçek bir kusur bulduğu tek örnektir.")


def section_8(doc: Document) -> None:
    h1(doc, "8. Ayrıntılı Karşılaştırma")

    h2(doc, "8.1 Teknik özellikler")
    table(doc,
          ["Konu", "Selenium", "Robot Framework", "Playwright"],
          [
              ["Mimari", "W3C WebDriver protokolü, ayrı sürücü süreci",
               "Selenium'un üzerinde keyword katmanı", "CDP/WebSocket ile doğrudan sürüş"],
              ["Otomatik bekleme", "Yok", "Yok (SeleniumLibrary)", "Var (actionability)"],
              ["Stale element", "Sık; elle yönetilir", "Sık; elle yönetilir", "Kavram olarak yok"],
              ["İzolasyon birimi", "Tarayıcı süreci", "Tarayıcı süreci", "BrowserContext"],
              ["Ağ katmanı erişimi", "Yok", "Yok", "Tam (route/abort/fulfill)"],
              ["API testi", "Ayrı kütüphane", "RequestsLibrary", "Dahili"],
              ["Mobil", "Appium ile", "Appium ile", "Emülasyon (gerçek cihaz yok)"],
              ["Tarayıcı desteği", "En geniş (IE dahil)", "Selenium ile aynı",
               "Chromium/Firefox/WebKit"],
              ["Sürücü yönetimi", "Selenium Manager (otomatik)", "Selenium Manager",
               "Kendi binary'lerini indirir"],
              ["Paralellik", "pytest-xdist (süreç)", "pabot (süreç)", "xdist + context (verimli)"],
          ],
          widths=[1.2, 1.7, 1.7, 1.8])

    h2(doc, "8.2 Ekip ve süreç")
    table(doc,
          ["Konu", "Selenium", "Robot Framework", "Playwright"],
          [
              ["Öğrenme eğrisi", "Orta–yüksek (Python + POM)", "Düşük (kod bilgisi gerekmez)",
               "Orta"],
              ["Programlama şart mı?", "Evet", "Hayır (temel testler için)", "Evet"],
              ["Kod okunabilirliği", "Geliştirici için iyi", "Herkes için iyi",
               "Geliştirici için çok iyi"],
              ["İş analisti katkısı", "Zor", "Kolay (CSV/Excel ile)", "Zor"],
              ["Rapor kalitesi", "Eklenti ile (pytest-html)", "Yerleşik ve çok güçlü",
               "Trace Viewer + eklenti"],
              ["Hata ayıklama", "Ekran görüntüsü + log", "log.html ağacı",
               "Trace Viewer (zaman içinde geri sarma)"],
              ["İş gücü bulunabilirliği", "En yüksek", "Orta", "Artıyor"],
              ["Topluluk/kaynak", "Devasa", "Orta", "Hızla büyüyor"],
          ],
          widths=[1.3, 1.7, 1.7, 1.7])

    h2(doc, "8.3 Aynı işin üç yazılışı")
    para(doc, "Bir ürünü sepete ekle ve rozetin güncellendiğini doğrula:", bold=True)

    para(doc, "Selenium + pytest", bold=True, color=ACCENT)
    code_block(doc,
               "def add_to_cart(self, product_name):\n"
               "    before = self.header.cart_count()\n"
               "    self.click(self._add_button(product_name))       # wait_clickable + retry\n"
               "    self.wait_visible(self._remove_button(product_name), timeout=10)\n"
               "    self.header.wait_cart_count(before + 1)\n"
               "    return self")

    para(doc, "Robot Framework", bold=True, color=ACCENT)
    code_block(doc,
               "Add Product To Cart\n"
               "    [Arguments]    ${product_name}\n"
               "    ${before} =    Get Cart Badge Count\n"
               "    ${add} =       Add To Cart Locator       ${product_name}\n"
               "    ${remove} =    Remove From Cart Locator  ${product_name}\n"
               "    Click And Wait For Element    ${add}    ${remove}    ${DEFAULT_TIMEOUT}\n"
               "    ${expected} =  Evaluate    ${before} + 1\n"
               "    Cart Badge Should Show    ${expected}")

    para(doc, "Playwright + pytest", bold=True, color=ACCENT)
    code_block(doc,
               "def add_to_cart(self, product_name):\n"
               "    before = self.header.cart_count()\n"
               "    self._add_button(product_name).click()            # auto-wait dahili\n"
               "    expect(self._remove_button(product_name)).to_be_visible()\n"
               "    self.header.expect_cart_count(before + 1)\n"
               "    return self")

    para(doc,
         "Üçü de aynı işi yapar ve benzer uzunluktadır. Fark, altta ne kadar kod "
         "bulunduğundadır: Selenium'un click() metodunun arkasında 60 satırlık bekleme ve "
         "yeniden deneme mantığı, Robot'un Click And Wait For Element keyword'ünün "
         "arkasında 40 satırlık etki doğrulama mantığı vardır. Playwright'ın click() "
         "metodunun arkasında framework'ün kendisi vardır.")

    h2(doc, "8.4 Maliyet analizi (kaba tahmin)")
    table(doc,
          ["Kalem", "Selenium", "Robot", "Playwright"],
          [
              ["İlk kurulum (altyapı)", "3–5 gün", "2–3 gün", "1–2 gün"],
              ["Test başına yazım", "Orta", "Hızlı (keyword varsa)", "Hızlı"],
              ["Bakım (UI değişikliği)", "Orta", "Orta", "Düşük (rol tabanlı locator)"],
              ["Kararsız test ayıklama", "Yüksek", "Yüksek", "Düşük"],
              ["CI süresi (100 test)", "≈10 dk", "≈15 dk", "≈3 dk"],
          ],
          widths=[1.9, 1.5, 1.5, 1.6])
    caption(doc,
            "CI süreleri bu projenin ölçümlerinden 100 teste ölçeklenerek tahmin edilmiştir; "
            "paralellik kapalıdır. Paralellikle üçü de kısalır ancak Playwright'ın avantajı "
            "context modeli sayesinde korunur.")


def section_9(doc: Document) -> None:
    h1(doc, "9. Uygulama Sırasında Karşılaşılan Gerçek Sorunlar")

    para(doc,
         "Bu bölüm, üç projeyi gerçekten çalıştırırken karşılaşılan ve çözülen sorunları "
         "anlatır. Karşılaştırmanın en değerli kısmı burasıdır: hangi framework'ün hangi "
         "sınıf sorunu ürettiğini teoriden değil, uygulamadan gösterir.")

    h2(doc, "Sorun 1 — Chrome 151, excludeSwitches ile açılmıyor (Selenium)")
    para(doc, "Belirti:", bold=True)
    code_block(doc, "SessionNotCreatedException: Chrome instance exited.")
    para(doc, "Kök neden:", bold=True)
    para(doc,
         "Birçok Selenium örneğinde önerilen "
         "options.add_experimental_option(\"excludeSwitches\", [\"enable-automation\"]) "
         "satırı, Chrome 151 ile tarayıcının açılır açılmaz kapanmasına yol açıyor. "
         "Seçenekler tek tek denenerek (bisection) bulundu.")
    para(doc, "Çözüm ve ders:", bold=True)
    para(doc,
         "Satır kaldırıldı. Bu, Selenium'un yapısal zayıflığının somut örneğidir: kod "
         "değişmez, tarayıcı güncellenir, paket toptan kırılır. Playwright'ta tarayıcı "
         "sürümü pakete kilitli olduğu için bu sınıf hatalar yaşanmaz.")

    h2(doc, "Sorun 2 — Menü animasyonu bitmeden okunan boş metinler (Selenium)")
    para(doc, "Belirti:", bold=True)
    code_block(doc, "Menu ogeleri: ['', '', 'Logout', 'Reset App State']   # ilk ikisi boş")
    para(doc, "Kök neden:", bold=True)
    para(doc,
         "Yan menü sağdan kayarak açılır. aria-hidden niteliği animasyon başlarken 'false' "
         "olur — yani \"açılmaya başladı\" demektir, \"açıldı\" değil. Selenium, görüntü "
         "alanı dışındaki elemanlar için .text değerini boş döner.")
    para(doc, "Çözüm:", bold=True)
    para(doc,
         "Tüm menü öğelerinin metni dolana kadar beklendi. Playwright'ta bu kontrol "
         "gerekmez; expect(...).to_be_visible() ve click(), elemanın konumunun "
         "sabitlenmesini zaten bekler.")

    h2(doc, "Sorun 3 — Sessizce yutulan tıklamalar (Robot Framework)")
    para(doc, "Belirti:", bold=True)
    para(doc,
         "Tıklama hata vermiyor, ancak menü açılmıyor. Test 15 saniye sonra alakasız bir "
         "\"element not visible\" hatasıyla düşüyor. Kök nedeni bulmak zor.")
    para(doc, "Ölçüm:", bold=True)
    para(doc,
         "İzole edilmiş denemelerde native tıklamanın yaklaşık %50 oranında etkisiz kaldığı, "
         "JavaScript tıklamasının ise 6/6 başarılı olduğu ölçüldü. SauceDemo'nun yan menüsü "
         "react-burger-menu bileşenidir ve tıklama alanı opacity: 0 olan bir buton "
         "katmanıdır; Chrome 151 + headless ortamında bu katmana giden fare olayları "
         "kayboluyor.")
    para(doc, "Çözüm:", bold=True)
    para(doc,
         "\"Etki doğrulamalı tıklama\" (Click Element Until) deseni kuruldu: önce gerçek "
         "tıklama denenir ve beklenen etkinin oluşup oluşmadığı doğrulanır; oluşmazsa "
         "JavaScript tıklamasına düşülür. Bu desen sonradan tüm etkileşimlere yayıldı.")
    code_block(doc,
               "Click Element Until\n"
               "    [Arguments]    ${locator}    @{verification}\n"
               "    ${ok} =    Run Keyword And Return Status\n"
               "    ...    Wait Until Keyword Succeeds    6s    500ms\n"
               "    ...    Click And Verify Effect    ${locator}    @{verification}\n"
               "    IF    not ${ok}\n"
               "        Click Element Via JavaScript    ${locator}\n"
               "        Run Keyword    @{verification}\n"
               "    END")
    para(doc, "Ders:", bold=True)
    para(doc,
         "\"Tıklandı mı?\" sorusunun cevabı güvenilir değildir; güvenilir olan tek ölçüt "
         "etkinin gerçekleşip gerçekleşmediğidir. Playwright bu kontrolü framework "
         "seviyesinde yapar; Selenium ve Robot'ta her projede yeniden kurulur.")

    h2(doc, "Sorun 4 — React kontrollü input gerçekten temizlenmiyor (Robot)")
    para(doc, "Belirti:", bold=True)
    para(doc,
         "\"Boş kullanıcı adı\" senaryosu, beklenen \"Username is required\" yerine "
         "\"do not match any user\" hatası aldı.")
    para(doc, "Kök neden:", bold=True)
    para(doc,
         "Clear Element Text (Selenium'un element.clear() metodu) DOM'daki value niteliğini "
         "boşaltır ama React'in iç state'ini güncellemez. Ekranda alan boş görünür, forma "
         "eski değer gider.")
    para(doc, "Çözüm:", bold=True)
    code_block(doc, "Press Keys    ${locator}    CTRL+a    DELETE   # gerçek klavye olayları")
    para(doc,
         "Bu tuzak Selenium'da da vardır; o projede her parametrik test yeni tarayıcı "
         "açtığı için hiç yüzeye çıkmadı. Playwright'ta locator.fill(\"\") doğru olayları "
         "ürettiği için görünmez olur.")

    h2(doc, "Sorun 5 — Vararg, global değişkeni gölgeledi (Robot)")
    para(doc, "Belirti:", bold=True)
    code_block(doc,
               "ValueError: Argument 'products' got value '[...]' (list)\n"
               "           that cannot be converted to dictionary.")
    para(doc, "Kök neden:", bold=True)
    para(doc,
         "Robot Framework'te değişken adları büyük/küçük harf, boşluk ve alt çizgi "
         "duyarsızdır. Bir keyword'ün @{products} vararg'ı, global ${PRODUCTS} fiyat "
         "sözlüğünü gölgeledi ve hesaplama keyword'ü sözlük yerine liste aldı.")
    para(doc, "Çözüm:", bold=True)
    para(doc,
         "Vararg @{product_names} olarak yeniden adlandırıldı. Bu, Robot'un okunabilirlik "
         "uğruna yaptığı bir tasarım tercihinin bedelidir; Python veya TypeScript'te "
         "böyle bir gölgeleme mümkün değildir.")

    h2(doc, "Sorun 6 — Headless'ta Maximize, pencereyi küçülttü (Robot)")
    para(doc, "Kök neden:", bold=True)
    para(doc,
         "Maximize Browser Window, headless modda pencereyi tarayıcının varsayılan sanal "
         "ekran boyutuna (800×600) düşürüyor ve --window-size=1920,1080 argümanını geri "
         "alıyor. Dar görüntü alanında bileşenler üst üste bindi.")
    para(doc, "Çözüm:", bold=True)
    code_block(doc,
               "IF    ${HEADLESS}\n"
               "    Set Window Size    ${WINDOW_WIDTH}    ${WINDOW_HEIGHT}\n"
               "ELSE\n"
               "    Run Keyword And Ignore Error    Maximize Browser Window\n"
               "END")

    h2(doc, "Sorun 7 — Uygulamanın kendi telemetri gürültüsü (Selenium + Playwright)")
    para(doc,
         "Konsol hatalarını izleyen test, SauceDemo'nun kendi telemetri servisine "
         "(backtrace.io) attığı ve her koşumda 401 dönen isteği yakaladı. Bu, test edilen "
         "uygulamanın bilinen davranışıdır; bilinen gürültü listesine alındı. Gerçek "
         "projelerde bu listenin kısa tutulması ve her maddesinin gerekçelendirilmesi "
         "gerekir — aksi halde test hiçbir şey yakalamaz hale gelir.")

    h2(doc, "Sorun 8 — Otomasyonun bulduğu gerçek kusur (Playwright)")
    para(doc,
         "Erişilebilirlik testi, SauceDemo'nun envanter sayfasında hiçbir başlık (h1–h6) "
         "elemanı olmadığını buldu; sayfa başlığı \"Products\" bir <span> olarak "
         "işlenmiş. Ekran okuyucu kullanıcıları sayfa yapısında gezinemez "
         "(WCAG 2.1 — 1.3.1). Test, xfail (bilinen açık kusur) olarak işaretlendi; "
         "kusur düzeltildiğinde test \"beklenmedik başarı\" verecek ve işaret kaldırılacak.",
         )
    para(doc,
         "Bu bulgu önemlidir: get_by_role() erişilebilirlik ağacını sorgular, CSS'i değil. "
         "Selenium ve Robot Framework'te bu bilgiye doğrudan erişimin bir yolu yoktur; "
         "aynı kusur o iki projede fark edilmeden kalırdı.", bold=True)

    h2(doc, "Sorun 9 — Kaynak baskısı altında toplu tıklama kaybı (Selenium tabanlı paketler)")
    para(doc, "Belirti:", bold=True)
    para(doc,
         "Robot paketi tek başına çalıştırıldığında tüm testler geçiyor; başka bir test "
         "paketi aynı anda çalışırken aynı paket testlerin 15'ini kaybediyor. Hataların "
         "tamamı \"tıklama etkisiz kaldı\" sınıfında. Aynı adımlar izole edilip tek başına "
         "denendiğinde ilk denemede çalışıyor.")
    para(doc, "Değerlendirme:", bold=True)
    para(doc,
         "WebDriver protokolü, her etkileşim için ayrı bir HTTP çağrısı ve ayrı bir sürücü "
         "süreci üzerinden çalışır. CPU baskısı altında bu zincirin zamanlaması bozuluyor ve "
         "fare olayları hedefe ulaşmadan kayboluyor. Playwright paketi aynı çekişme altında "
         "etkilenmedi; tek bir WebSocket bağlantısı üzerinden çalışması ve her etkileşim "
         "öncesi elemanın gerçekten tıklanabilir olduğunu doğrulaması bunu açıklıyor.")
    para(doc, "Ders:", bold=True)
    para(doc,
         "Bu koşul CI sunucularında kural, istisnadır değil: paylaşımlı runner'lar, paralel "
         "işler, sınırlı CPU. \"Yerelde geçiyor, CI'da bazen kırmızı\" şikâyetinin en sık "
         "sebeplerinden biri budur. Framework seçiminin kararlılığa etkisi, en net biçimde "
         "burada görülür.", bold=True)


def section_10(doc: Document) -> None:
    h1(doc, "10. Hangi Durumda Hangisi? — Karar Rehberi")

    h2(doc, "Playwright seçin, eğer…")
    bullets(doc, [
        "Yeni bir projeye sıfırdan başlıyorsanız.",
        "Uygulamanız modern bir SPA ise (React, Vue, Angular).",
        "CI süresi ve test kararlılığı önceliğinizse.",
        "Ağ katmanını taklit etmeniz gerekiyorsa (backend hazır değil, hata senaryoları).",
        "API ve UI testlerini aynı çatı altında istiyorsanız.",
        "Ekibiniz geliştirici ağırlıklıysa.",
    ])

    h2(doc, "Robot Framework seçin, eğer…")
    bullets(doc, [
        "Ekipte programlama bilmeyen üyeler test yazacaksa.",
        "Kabul testleri (acceptance test) müşteriyle birlikte yazılacaksa.",
        "Rapor kurum içinde kanıt olarak kullanılacaksa (regülasyon, denetim).",
        "Test senaryoları Excel/CSV'de yönetiliyorsa.",
        "Yalnızca web değil, veritabanı/SSH/API/mobil aynı çatıda test edilecekse.",
    ])
    para(doc,
         "Öneri: Robot Framework seçiyorsanız SeleniumLibrary yerine Browser kütüphanesini "
         "(Playwright tabanlı) değerlendirin. Robot'un okunabilirliği ile Playwright'ın "
         "kararlılığı birleşir; bu projede yaşanan tıklama sorunlarının çoğu ortadan kalkar.",
         bold=True)

    h2(doc, "Selenium seçin, eğer…")
    bullets(doc, [
        "Kurumda zaten büyük bir Selenium yatırımı varsa.",
        "Çok geniş bir tarayıcı/sürüm matrisi desteklemeniz gerekiyorsa.",
        "Selenium Grid veya BrowserStack/Sauce Labs altyapısı kuruluysa.",
        "Appium ile mobil otomasyonu aynı bilgi birikimiyle yürütmek istiyorsanız.",
        "İş gücü bulunabilirliği kritik bir kısıtsa.",
    ])

    h2(doc, "Geçiş stratejisi")
    para(doc,
         "Mevcut bir Selenium paketi Playwright'a taşınacaksa, tamamını bir kerede "
         "dönüştürmek yerine şu sıra önerilir: (1) yeni testleri Playwright ile yazın, "
         "(2) en çok kararsızlık üreten mevcut testleri taşıyın, (3) Page Object "
         "katmanını taşıyın — testler büyük ölçüde aynı kalır, (4) Selenium paketini "
         "yalnızca gerçekten gereken tarayıcılar için tutun.")


def section_11(doc: Document) -> None:
    h1(doc, "11. Sonuç ve Öneri")

    para(doc,
         "Üç framework de aynı işi yapabildi. Üçüyle de aynı senaryolar yazıldı, üçü de "
         "gerçek hatalar yakaladı, üçü de bakımı yapılabilir bir mimariye oturdu. "
         "Dolayısıyla soru \"hangisi çalışır\" değil, \"hangisi hangi maliyetle çalışır\".")

    h2(doc, "Üç cümlelik özet")
    para(doc,
         "Selenium en yaygın ve en esnek olandır; karşılığında kararlılık altyapısını "
         "siz yazarsınız. Robot Framework en okunabilir ve en iyi raporlayan olandır; "
         "karşılığında karmaşık mantık için Python'a inmeniz gerekir. Playwright en hızlı "
         "ve en kararlı olandır; karşılığında ekosistem genişliğinden ve bazı tarayıcı "
         "desteklerinden feragat edersiniz.")

    h2(doc, "Bu projenin en net bulgusu")
    para(doc,
         "Framework seçiminin en büyük etkisi, kod miktarında değil kararsız test "
         "(flaky test) oranındadır. Selenium ve Robot projelerinde ilk koşumda çıkan "
         "hataların çoğunluğu bekleme ve tıklama kararlılığıyla ilgiliydi ve çözümleri "
         "toplam 100+ satır ek altyapı gerektirdi. Playwright projesinde bu sınıftan tek "
         "bir hata çıkmadı. Uzun vadede bir test paketinin gerçek maliyeti, yazılırken "
         "değil, her hafta \"neden bu test bazen kırmızı yanıyor\" sorusuna harcanan "
         "zamanda birikir.",
         bold=True)

    h2(doc, "Nihai öneri")
    table(doc,
          ["Durum", "Öneri"],
          [
              ["Yeni web projesi, geliştirici ağırlıklı ekip", "Playwright (Python veya TypeScript)"],
              ["Karma ekip, iş analisti katkısı, kabul testleri", "Robot Framework + Browser kütüphanesi"],
              ["Mevcut Selenium yatırımı, geniş tarayıcı matrisi", "Selenium (kademeli geçiş planıyla)"],
              ["Mobil + web birlikte", "Selenium/Appium veya Robot + AppiumLibrary"],
              ["Regülasyonlu sektör, denetim kanıtı gereksinimi", "Robot Framework"],
          ],
          widths=[3.2, 3.2])

    doc.add_paragraph()
    para(doc,
         "Bu raporun dayandığı üç proje, aynı depoda çalışır durumdadır. Her biri kendi "
         "README dosyasında ayrıntılı olarak belgelenmiştir ve run_tests.ps1 betiğiyle "
         "doğrudan çalıştırılabilir.", italic=True, color=MUTED)


def appendix_a(doc: Document) -> None:
    h1(doc, "Ek A — Çalıştırma Komutları")

    h2(doc, "Kurulum")
    code_block(doc,
               "python -m venv .venv\n"
               ".\\.venv\\Scripts\\Activate.ps1\n"
               "\n"
               "pip install -r 01-selenium-pytest\\requirements.txt\n"
               "pip install -r 02-robot-framework\\requirements.txt\n"
               "pip install -r 03-playwright-pytest\\requirements.txt\n"
               "\n"
               "python -m playwright install chromium firefox webkit")

    h2(doc, "Selenium")
    code_block(doc,
               "cd 01-selenium-pytest\n"
               ".\\run_tests.ps1                    # tümü, headless\n"
               ".\\run_tests.ps1 -Marker smoke      # sadece smoke\n"
               ".\\run_tests.ps1 -Headed            # tarayıcı görünür\n"
               ".\\run_tests.ps1 -Parallel 4        # 4 süreç paralel\n"
               "python -m pytest -m \"cart or checkout\" -v")

    h2(doc, "Robot Framework")
    code_block(doc,
               "cd 02-robot-framework\n"
               ".\\run_tests.ps1                    # tümü\n"
               ".\\run_tests.ps1 -Tags smoke\n"
               ".\\run_tests.ps1 -Suite 03_cart     # tek suite\n"
               ".\\run_tests.ps1 -Parallel 4        # pabot ile\n"
               "\n"
               "robot --outputdir results \\\n"
               "      --variablefile variables/environment.py:prod:chrome:true tests/")

    h2(doc, "Playwright")
    code_block(doc,
               "cd 03-playwright-pytest\n"
               ".\\run_tests.ps1                    # tümü\n"
               ".\\run_tests.ps1 -AllBrowsers       # chromium + firefox + webkit\n"
               ".\\run_tests.ps1 -Debug             # inspector + yavaşlatılmış\n"
               "\n"
               "python -m playwright show-trace reports\\artifacts\\<test>\\trace.zip\n"
               "python -m playwright codegen https://www.saucedemo.com")

    h2(doc, "Bu raporu yeniden üretmek")
    code_block(doc, "pip install python-docx\npython docs\\generate_report.py")


def appendix_b(doc: Document) -> None:
    h1(doc, "Ek B — Dosya Yapısı")

    code_block(doc,
               "Automation/\n"
               "├── README.md\n"
               "├── Otomasyon_Karsilastirma_Raporu.docx    <- bu rapor\n"
               "├── shared/\n"
               "│   ├── users.json                          8 login senaryosu\n"
               "│   └── products.json                       6 ürün + KDV + müşteri\n"
               "├── docs/generate_report.py                 rapor üreteci\n"
               "│\n"
               "├── 01-selenium-pytest/                     27 dosya · 2.748 satır · 61 test\n"
               "│   ├── conftest.py                         fixture + hook (201 satır)\n"
               "│   ├── pytest.ini\n"
               "│   ├── config/{config.yaml, settings.py}\n"
               "│   ├── core/\n"
               "│   │   ├── driver_factory.py               Chrome/Firefox/Edge\n"
               "│   │   ├── base_page.py                    288 satır\n"
               "│   │   ├── waits.py                        127 satır, 7 özel koşul\n"
               "│   │   ├── soft_assert.py\n"
               "│   │   ├── exceptions.py\n"
               "│   │   └── logger.py\n"
               "│   ├── pages/{login, inventory, cart, checkout, product_detail}\n"
               "│   │   └── components/header_component.py\n"
               "│   ├── utils/{parsers, data_loader}\n"
               "│   └── tests/                              6 dosya\n"
               "│\n"
               "├── 02-robot-framework/                     14 dosya · 2.068 satır · 51 test\n"
               "│   ├── tests/\n"
               "│   │   ├── __init__.robot                  suite setup/teardown\n"
               "│   │   ├── 01_login.robot                  [Template] veri odaklı\n"
               "│   │   ├── 01b_login_datadriven.robot      DataDriver + CSV\n"
               "│   │   ├── 02_inventory.robot\n"
               "│   │   ├── 03_cart.robot\n"
               "│   │   ├── 04_checkout.robot\n"
               "│   │   └── 05_e2e.robot\n"
               "│   ├── resources/\n"
               "│   │   ├── common.resource                 335 satır\n"
               "│   │   └── pages/{login, inventory, cart, checkout}.resource\n"
               "│   ├── libraries/SauceDemoLibrary.py       200 satır Python keyword\n"
               "│   ├── variables/environment.py            ortam değişkenleri\n"
               "│   └── data/login_scenarios.csv\n"
               "│\n"
               "└── 03-playwright-pytest/                   27 dosya · 2.361 satır · 75 test\n"
               "    ├── conftest.py                         222 satır (%60'ı açıklama)\n"
               "    ├── pytest.ini\n"
               "    ├── config/settings.py\n"
               "    ├── core/base_page.py                   109 satır\n"
               "    ├── pages/{login, inventory, cart, checkout, product_detail}\n"
               "    │   └── components/header_component.py\n"
               "    ├── utils/{parsers, data_loader}\n"
               "    └── tests/\n"
               "        ├── ui/                             Selenium & Robot ile aynı 5 dosya\n"
               "        └── advanced/                       diğerlerinde karşılığı olmayanlar\n"
               "            ├── test_network_interception.py\n"
               "            ├── test_visual_and_devices.py\n"
               "            └── test_api.py")


def main() -> None:
    doc = Document()
    setup_styles(doc)

    build_cover(doc)
    build_toc(doc)
    section_1(doc)
    section_2(doc)
    section_3(doc)
    section_4(doc)
    section_5(doc)
    section_6(doc)
    section_7(doc)
    section_8(doc)
    section_9(doc)
    section_10(doc)
    section_11(doc)
    appendix_a(doc)
    appendix_b(doc)

    doc.save(OUTPUT)
    print(f"Rapor olusturuldu: {OUTPUT}")
    print(f"Bolum sayisi: 11 + 2 ek")


if __name__ == "__main__":
    main()
